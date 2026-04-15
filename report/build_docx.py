"""Assemble the final-year academic report (Microsoft Word)."""
from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
ASSETS = Path(__file__).resolve().parent / "assets"
METRICS = ROOT / "models" / "evaluation_metrics.json"


class StreamlitProjectReport:
    """Builds the printable submission document with embedded diagrams."""

    def __init__(self) -> None:
        self.doc = Document()
        self.metrics = json.loads(METRICS.read_text(encoding="utf-8")) if METRICS.exists() else {}

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)

    def build(self) -> None:
        self._title()
        self.doc.add_page_break()
        self._front_matter()
        self.doc.add_page_break()
        self._design()
        self.doc.add_page_break()
        self._ml_section()
        self.doc.add_page_break()
        self._implementation_results()
        self.doc.add_page_break()
        self._closing()

    def _title(self) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Smart Undergraduate Course Navigator\n"
            "Intelligent Decision Support for Course Selection and Academic Performance Optimization at UMaT"
        )
        run.bold = True
        sub = self.doc.add_paragraph("Final Year Project - Streamlit Implementation")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = self.doc.add_paragraph(
            f"University of Mines and Technology (UMaT)\n{date.today():%B %d, %Y}\n"
            "GitHub (placeholder): https://github.com/dogahwisdom/smart-course-navigator"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _front_matter(self) -> None:
        for heading, body in [
            (
                "Abstract",
                "This artefact documents a Streamlit-based decision support system that couples cohort analytics "
                "with supervised machine learning to forecast pass probabilities, recommend semester loads, and "
                "classify academic risk. Three classifiers-logistic regression, decision trees, and random forests-"
                "are trained on a realistic synthetic engineering dataset (≥1000 rows) spanning mining, geological, "
                "electrical, and mechanical programs. The best-performing model is serialized with joblib and invoked "
                "from a multi-page Streamlit experience built entirely in Python.",
            ),
            (
                "Introduction",
                "Engineering undergraduates must balance accreditation requirements, laboratory intensity, and "
                "individual preparedness. The navigator operationalizes historical signals into individualized guidance.",
            ),
            (
                "Problem Statement",
                "Students lack transparent, quantitative aids when sequencing demanding courses, while "
                "administrators need reproducible pipelines for refreshing analytics.",
            ),
        ]:
            self.doc.add_heading(heading, level=1)
            self.doc.add_paragraph(body)

        self.doc.add_heading("Objectives", level=1)
        for item in [
            "Deliver Streamlit dashboards for analytics, prediction, recommendations, and model governance.",
            "Train and compare three ML models with accuracy, precision, recall, and F1-score.",
            "Expose reasoning strings aligned with cohort behaviour and modeled probabilities.",
        ]:
            self.doc.add_paragraph(item, style="List Bullet")

        self.doc.add_heading("Literature Review", level=1)
        self.doc.add_paragraph(
            "Decision support systems integrate data, models, and interfaces for semi-structured decisions (Power, "
            "Sharda, & Burstein, 2015). Educational data mining applies classification to academic traces (Romero & "
            "Ventura, 2013). Responsible deployment foregrounds transparency and student agency (Williamson, 2018)."
        )

        self.doc.add_heading("Methodology", level=1)
        self.doc.add_paragraph(
            "Synthetic data emulate registrar-style fields; preprocessing scales numerics and one-hot encodes "
            "categoricals; class imbalance is mitigated via class_weight within estimators; evaluation uses a "
            "stratified hold-out split."
        )

    def _design(self) -> None:
        self.doc.add_heading("System Design", level=1)
        self.doc.add_paragraph(
            "The Streamlit runtime orchestrates UI, calls utility modules for analytics, recommendations, and risk "
            "analysis, and loads the persisted sklearn pipeline from disk."
        )
        for name in [
            "fig_architecture.png",
            "fig_ml_workflow.png",
        ]:
            p = ASSETS / name
            if p.exists():
                self.doc.add_picture(str(p), width=Inches(5.9))
        self.doc.add_heading("Dataset Description", level=1)
        self.doc.add_paragraph(
            "Fields include student_id, program, GPA, course identifiers, difficulty, attempts, attendance, "
            "credit load, grade, and pass/fail labels for supervised learning."
        )

    def _ml_section(self) -> None:
        self.doc.add_heading("Machine Learning Approach", level=1)
        if self.metrics:
            self.doc.add_paragraph(
                f"Selected model: {self.metrics.get('selected_model')}. "
                f"{self.metrics.get('selection_criterion', '')}"
            )
            tbl = self.doc.add_table(rows=1, cols=5)
            hdr = tbl.rows[0].cells
            hdr[0].text = "Model"
            hdr[1].text = "Accuracy"
            hdr[2].text = "Precision"
            hdr[3].text = "Recall"
            hdr[4].text = "F1"
            for row in self.metrics.get("results", []):
                cells = tbl.add_row().cells
                cells[0].text = row["model"]
                cells[1].text = str(row["accuracy"])
                cells[2].text = str(row["precision"])
                cells[3].text = str(row["recall"])
                cells[4].text = str(row["f1_score"])
        for name in ["fig_model_compare.png", "fig_feature_importance.png"]:
            p = ASSETS / name
            if p.exists():
                self.doc.add_picture(str(p), width=Inches(5.9))

    def _implementation_results(self) -> None:
        self.doc.add_heading("System Implementation (Streamlit)", level=1)
        self.doc.add_paragraph(
            "The implementation follows the prescribed structure: app.py for the home page, pages/ for modular "
            "views, utils/ for preprocessing and ML services, and notebooks/training.ipynb for reproducible training."
        )
        ui = ASSETS / "fig_streamlit_ui.png"
        if ui.exists():
            self.doc.add_picture(str(ui), width=Inches(5.8))
        self.doc.add_heading("Results and Evaluation", level=1)
        self.doc.add_paragraph(
            "Quantitative metrics are summarized in the performance dashboard and embedded figures. Qualitative "
            "review confirms that reasoning strings map to cohort statistics and individualized probabilities."
        )
        self.doc.add_heading("Discussion", level=1)
        self.doc.add_paragraph(
            "Synthetic data enable safe academic demonstration; production deployment requires governed extracts, "
            "bias audits, and integration with official advising workflows."
        )

    def _closing(self) -> None:
        self.doc.add_heading("Conclusion", level=1)
        self.doc.add_paragraph(
            "The Streamlit navigator satisfies the functional brief: prediction, analytics, recommendations, risk "
            "tiers, and model transparency within a single Python codebase."
        )
        self.doc.add_heading("Recommendations", level=1)
        self.doc.add_paragraph(
            "Adopt institutional authentication, connect to relational databases, and schedule quarterly model "
            "retraining with drift monitoring."
        )
        self.doc.add_heading("References (APA-style samples)", level=1)
        for ref in [
            "Power, D. J., Sharda, R., & Burstein, F. (2015). Decision support systems (3rd ed.). Wiley.",
            "Romero, C., & Ventura, S. (2013). Data mining in education. Wiley Interdisciplinary Reviews.",
            "Pedregosa, F. et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        ]:
            self.doc.add_paragraph(ref, style="List Bullet")

        self.doc.add_heading("Contributors", level=1)
        team = [
            ("Wisdom Dogah", "FCM.41.020.099.23"),
            ("Faustian Nyamekye", "FCM.41.020.146.23"),
            ("Owusu Appiah Barimah Kofi Duodu", "FCM.41.020.159.23"),
            ("De-Graft Prince Kweku", "F.C.M. 41.020.098.23"),
            ("Cudjoe Jennifer Abena", "FCM. 41.020.096.23"),
            ("Danquah Joseph", "FCM.41.020.097.23"),
            ("Kelvin Kandibiga", "FCM.41.020.118.23"),
        ]
        for name, index_no in team:
            self.doc.add_paragraph(f"{name} - {index_no}")

        self.doc.add_heading("Appendices", level=1)
        self.doc.add_paragraph(
            "Appendix A: Directory tree (app.py, pages/, utils/, data/, models/). Appendix B: API-free architecture "
            "because Streamlit embeds server logic. Appendix C: evaluation_metrics.json stores raw metrics for audits."
        )


def main() -> None:
    report = StreamlitProjectReport()
    report.build()
    out = ROOT / "report" / "Smart_Course_Navigator_Streamlit_Report.docx"
    report.save(out)
    print("Wrote", out)


if __name__ == "__main__":
    main()
